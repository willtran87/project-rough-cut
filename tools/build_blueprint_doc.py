from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "docs" / "Rough_Cut_Game_Blueprint.docx"
HERO_IMAGE = (
    PROJECT_ROOT
    / "assets"
    / "characters"
    / "joe"
    / "source_reference"
    / "portrait-inspired-man_mow-walk-chaotic-head_FRONT_v2_keyframes.png"
)

FONT = "Calibri"
DARK_GREEN = RGBColor(31, 77, 58)
MID_GREEN = RGBColor(46, 110, 80)
DEEP_BROWN = RGBColor(92, 58, 46)
MOWER_ORANGE = RGBColor(198, 74, 27)
INK = RGBColor(32, 38, 35)
MUTED = RGBColor(96, 105, 100)
LIGHT_GREEN = "E3EAE4"
LIGHT_ORANGE = "F7E7DF"
LIGHT_GRAY = "F2F4F3"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="B8C3BC", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def style_table_text(table, header=True):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9.2 if row_index > 0 else 9.4,
                        color=INK,
                        bold=bool(header and row_index == 0),
                    )
        if header and row_index == 0:
            repeat_header_row(row)
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GREEN)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(paragraph)
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True, color=INK)
        remainder = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(remainder, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=INK)
    return paragraph


def add_bullet(doc, text, level=0):
    style_name = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    set_run_font(run, color=INK)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run_font(run, color=INK)
    return paragraph


def add_callout(doc, label, text, fill=LIGHT_ORANGE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.line_spacing = 1.12
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "D8B7A6")
        borders.append(border)
    p_pr.append(borders)
    lead = paragraph.add_run(f"{label}: ")
    set_run_font(lead, size=10.2, bold=True, color=MOWER_ORANGE)
    body = paragraph.add_run(text)
    set_run_font(body, size=10.2, color=INK)


def add_data_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_values in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_values):
            row_cells[index].text = str(value)
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    style_table_text(table)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, DARK_GREEN, 18, 10),
        "Heading 2": (13, MID_GREEN, 14, 7),
        "Heading 3": (12, DEEP_BROWN, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(11)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.375)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.188)
    styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.625)
    styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.188)
    styles["List Number"].paragraph_format.left_indent = Inches(0.375)
    styles["List Number"].paragraph_format.first_line_indent = Inches(-0.188)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.paragraph_format.space_before = Pt(0)
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5),
        alignment=WD_TAB_ALIGNMENT.RIGHT,
    )
    left_run = header_paragraph.add_run("ROUGH CUT")
    set_run_font(left_run, size=8.2, bold=True, color=MUTED)
    tab_run = header_paragraph.add_run("\t")
    set_run_font(tab_run, size=8.2, color=MUTED)
    right_run = header_paragraph.add_run("GAME BLUEPRINT")
    set_run_font(right_run, size=8.2, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = footer.add_run("Project blueprint  |  ")
    set_run_font(label, size=8.5, color=MUTED)
    add_page_number(footer)


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(38)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("GAME CONCEPT & VERTICAL SLICE BLUEPRINT")
    set_run_font(run, size=10, bold=True, color=MOWER_ORANGE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("ROUGH CUT")
    set_run_font(run, size=30, bold=True, color=DARK_GREEN)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    run = subtitle.add_run("A JOE HORROR GAME")
    set_run_font(run, size=15, bold=True, color=DEEP_BROWN)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_after = Pt(22)
    run = tagline.add_run("The course closes at dusk. Joe does not.")
    set_run_font(run, size=11.5, italic=True, color=MUTED)

    if HERO_IMAGE.exists():
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.space_after = Pt(18)
        image_run = image_paragraph.add_run()
        shape = image_run.add_picture(str(HERO_IMAGE), width=Inches(5.85))
        shape._inline.docPr.set(
            "descr",
            "Ten-frame pixel-art animation board of Joe pushing a red lawn mower with erratic arm and head movement.",
        )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    run = meta.add_run("FIRST-PERSON 2.5D PIXEL HORROR  |  DESIGN STATUS: CONCEPT")
    set_run_font(run, size=9.5, bold=True, color=MID_GREEN)

    doc.add_page_break()


def add_snapshot(doc):
    add_heading(doc, "1. Project vision", 1)
    add_callout(
        doc,
        "Logline",
        "After an insurance-company golf outing becomes Joe's unauthorized 'course optimization pilot,' a stranded coworker must cross nine increasingly hostile holes while the software Product Owner—obsessed with grass, golf, and measurable outcomes—mows away hiding places and treats the player as unplanned scope.",
    )
    add_data_table(
        doc,
        ["Field", "Working definition"],
        [
            ("Genre", "First-person survival horror with stealth, pursuit, environmental manipulation, and deadpan comedy."),
            ("Format", "Low-resolution 3D course geometry with high-detail billboard sprites and nearest-neighbor presentation."),
            ("Player fantasy", "Outthink an apparently chaotic pursuer by reading the course, controlling sound, and exploiting momentum."),
            ("Narrative setup", "A company golf outing runs past dusk. Joe takes control of the course's maintenance systems to finish an obsessive optimization plan, trapping one coworker inside."),
            ("Target session", "A 2-3 hour first campaign; 10-20 minute replayable hole scenarios."),
            ("Primary antagonist", "Joe, a software Product Owner employed by an insurance company—not an adjuster—whose grass-and-golf obsession turns backlog ownership, sprint goals, acceptance criteria, and a lawn mower into a terrifying pursuit system."),
            ("Recommended engine", "Godot for a lightweight 3D world, billboard sprites, data-driven AI, and rapid vertical-slice iteration."),
        ],
        [2700, 6660],
    )

    add_heading(doc, "Experience goals", 2)
    for item in (
        "Sustain anticipation longer than direct pursuit.",
        "Make every chase alter the course by permanently cutting cover.",
        "Let players survive through observation and improvisation rather than combat.",
        "Use absurd golf-course normality to release tension without making Joe harmless.",
        "Create streamer-readable situations with clear cause, consequence, and recovery.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Design pillars", 2)
    pillars = [
        ("Readable chaos", "Joe looks erratic, but head, hand, engine, and posture cues reveal his actual state."),
        ("The course remembers", "Mowing, opened shortcuts, activated sprinklers, and displaced props persist during a run."),
        ("Sound is gameplay", "The mower engine, grass movement, golf impacts, and PA system are navigation and deception tools."),
        ("Comedy stays diegetic", "The course behaves professionally and politely while something impossible is happening."),
        ("Short pursuit, long dread", "Chases are intense and decisive; anticipation, investigation, and aftermath occupy more time."),
        ("Mastery compounds", "Recoveries, bait plays, progress, and clean contact breaks can be linked into a short Delivery Chain that rewards expressive play without reducing danger."),
    ]
    for label, detail in pillars:
        add_body(doc, f"{label}. {detail}", bold_lead=f"{label}.")


def add_format_and_loop(doc):
    add_heading(doc, "2. Recommended game format", 1)
    add_body(
        doc,
        "Use a first-person 2.5D presentation: low-poly golf-course geometry and pixelated textures rendered at 320x180 or 640x360 internal resolution, with Joe displayed as a detailed animated billboard. Joe's tendency to face the player becomes uncanny rather than technically limiting.",
    )
    add_heading(doc, "Art-direction rules", 2)
    for item in (
        "Keep the course readable: open sightlines, clear silhouettes, and selective fog instead of uniform darkness.",
        "Render Joe at greater detail than the environment so he feels distressingly real inside a crude simulation.",
        "Animate Joe at 8-12 FPS while the camera remains smooth.",
        "Use sickly fairway greens, desaturated twilight blues, sand beige, deep tree-line blacks, and mower orange-red.",
        "Show persistent mowing strips, bent grass, sprinkler particles, and fog layers as gameplay information.",
        "Avoid heavy VHS overlays that obscure hazards, interaction points, or Joe's telegraphs.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Core loop", 2)
    core_steps = (
        "Enter a new hole and identify the exit condition.",
        "Read the terrain, mower sound, and available tools.",
        "Complete two or three noisy objectives while managing concealment.",
        "Redirect or evade Joe as he permanently changes the route.",
        "Link smart plays inside the Delivery Chain window for optional score mastery.",
        "Attempt an optional golf challenge for an upgrade or story fragment.",
        "Reach the next tee, maintenance passage, or clubhouse checkpoint.",
    )
    for step in core_steps:
        add_number(doc, step)

    add_callout(
        doc,
        "Pacing target",
        "Approximately 70 percent anticipation and navigation, 20 percent active pursuit, and 10 percent recovery, comedy, and narrative discovery.",
        fill=LIGHT_GREEN,
    )


def add_joe_system(doc):
    add_heading(doc, "3. Joe behavior system", 1)
    add_body(
        doc,
        "Joe should not be truly random. His animation is theatrical noise layered over a deterministic state machine. Players who study him gain meaningful predictive power.",
    )
    add_data_table(
        doc,
        ["State", "Visible or audible cue", "Joe's behavior", "Player response"],
        [
            ("Mowing", "Steady engine and two-hand grip", "Follows a mowing route and removes tall grass.", "Stay hidden, reroute, or exploit his predictable line."),
            ("Listening", "Engine idles; head snaps toward a sound", "Stops advancing and samples recent noise events.", "Freeze, move under ambient noise, or create a second distraction."),
            ("Investigating", "Uneven throttle and slow head tracking", "Moves toward disturbed grass, impacts, or activated devices.", "Leave before he reaches the source; prepare a sharp-turn escape."),
            ("Watching", "Motionless at long range, directly facing player", "Builds dread and blocks a route without immediate attack.", "Break line of sight and question whether he truly saw you."),
            ("Charging", "Head lowers, grip stabilizes, engine pitch rises", "Accelerates rapidly in a straight line.", "Commit to a late lateral dodge around solid terrain."),
            ("Unstable pursuit", "One arm releases; shoulders and head jerk", "Swerves and corrects course unpredictably within limits.", "Use trees, signs, bunker lips, and narrow route changes."),
            ("Silent search", "Engine cuts out; mower remains unattended", "Joe searches on foot or repositions out of sight.", "Treat silence as danger; use grass movement and environmental audio."),
            ("Course correction", "Distant restart or impossible engine angle", "Uses a shortcut or previous player route to regain pressure.", "Avoid repeating the same escape path."),
        ],
        [1450, 2300, 2900, 2710],
    )

    add_heading(doc, "Animation-to-behavior language", 2)
    cues = (
        "Both hands steady: a straight, committed movement is coming.",
        "One hand off the mower: expect wobble, a turn, or a delayed correction.",
        "Head snaps toward a location: Joe registered a sound there.",
        "Head lowered: acceleration or direct pursuit is imminent.",
        "Head tilted while stationary: Joe is listening, not merely idling.",
        "Sudden normal posture: use as a rare prelude to a dangerous phase transition.",
    )
    for cue in cues:
        add_bullet(doc, cue)


def add_player_course(doc):
    add_heading(doc, "4. Player and course mechanics", 1)
    add_heading(doc, "Player verbs", 2)
    player_verbs = [
        ("Walk, crouch, sprint", "Movement changes noise, visibility, and stamina."),
        ("Hide in rough", "Tall grass conceals but slows movement and records disturbance."),
        ("Throw golf balls", "Creates discrete sound events that can redirect Joe."),
        ("Strike objects", "A club can ring signs, pipes, bells, or balls; it is not a primary weapon."),
        ("Operate sprinklers", "Masks audio, reveals silhouettes, and makes hard surfaces slick."),
        ("Drive a golf cart", "Fast, loud, fuel-limited traversal with awkward handling."),
        ("Use maintenance access", "Keys and switches open sheds, tunnels, and shortcuts."),
        ("Attempt golf shots", "Optional high-noise challenges grant upgrades and narrative rewards."),
    ]
    for label, detail in player_verbs:
        add_body(doc, f"{label}. {detail}", bold_lead=f"{label}.")

    add_heading(doc, "Terrain rules", 2)
    add_data_table(
        doc,
        ["Terrain", "Player effect", "Joe effect"],
        [
            ("Fairway", "Fast, exposed movement and clear sightlines.", "High mower speed and reliable charges."),
            ("Tall rough", "Concealment, noise, and reduced speed.", "Can be permanently removed by mowing."),
            ("Bunker", "Slow movement with highly visible footprints.", "Mower loses momentum or must route around."),
            ("Pond or stream", "Hard boundary, reflection cue, and forced detour.", "Constrains approach angles."),
            ("Tree line", "Line-of-sight break and sharp-turn protection.", "Creates collision risk during a charge."),
            ("Cart path", "Fast walking and driving but little cover.", "Supports aggressive course correction."),
            ("Maintenance area", "Tools, keys, fuel, and temporary shelter.", "High-value investigation destination."),
        ],
        [1700, 3830, 3830],
    )

    add_heading(doc, "Persistent course state", 2)
    for item in (
        "Mowed grass remains removed for the rest of the run.",
        "Opened gates and maintenance shortcuts remain available.",
        "Activated sprinklers cycle on timers instead of resetting immediately.",
        "Moved carts, rakes, signs, and barriers remain where the player leaves them.",
        "Joe becomes faster on routes he has already cleared.",
    ):
        add_bullet(doc, item)


def add_tone_audio(doc):
    add_heading(doc, "5. Suspense, horror, drama, and comedy", 1)
    add_heading(doc, "Tension cycle", 2)
    for step in (
        "Unease: evidence of Joe appears before he does.",
        "Uncertainty: location, route, or intention is unclear.",
        "Recognition: an animation or sound reveals his state.",
        "Pursuit: the player commits to a terrain-based escape.",
        "Escape: the player reaches temporary safety or redirects Joe.",
        "Release: a mundane golf-course detail creates a brief laugh.",
        "Complication: the next objective removes certainty or cover.",
    ):
        add_number(doc, step)

    add_heading(doc, "Audio direction", 2)
    for item in (
        "Use the mower as the primary musical instrument: idle, load, acceleration, distance filtering, occlusion, restart, and mechanical failure.",
        "Make engine silence more threatening than constant volume.",
        "Allow hills, trees, sheds, and fog zones to distort perceived direction.",
        "Use sprinklers, insects, wind, and public-address announcements as masking layers.",
        "Reserve conventional music for title, safe-room, revelation, and endgame beats.",
        "Give the golf cart a cheerful reverse alarm that becomes tactically dangerous.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Comedy rules", 2)
    add_body(
        doc,
        "Joe never acknowledges that the situation is funny. Humor comes from his absolute sincerity: software-product, Agile, Scrum, and delivery language applied to mortal danger, course etiquette, machinery, and the player's desperate use of ordinary golf equipment. Insurance is his industry context, not his profession.",
    )
    comedy_examples = (
        "A speaker announces, 'Your current route is outside agreed scope,' as Joe charges.",
        "Joe shouts, 'Acceptance criteria require a uniform cut height!' while mowing through the player's cover.",
        "A failure screen reads, 'SPRINT TERMINATED: failed Joe's acceptance review.'",
        "A course map has been converted into a product roadmap with every hiding place marked as technical debt.",
        "Joe pauses to repair a divot and update a clipboard, then resumes the chase without comment.",
    )
    for example in comedy_examples:
        add_bullet(doc, example)

    add_heading(doc, "Joe's backstory", 2)
    add_body(
        doc,
        "Joe is a software Product Owner employed by an insurance company. He is not an adjuster, does not investigate losses, and does not decide claims. He owns value, ordering, and clarity for an internal software backlog: translating stakeholder ambiguity into product goals, prioritized user stories, acceptance criteria, release increments, dependency decisions, and measurable outcomes. Away from work he applies that same intensity to grass and golf: mowing patterns are roadmaps, divots are defects, bunkers are blockers, weather is a dependency, and an imperfect fairway is an increment that cannot be accepted.",
    )
    for item in (
        "Joe began in the insurer's software division by owning a policy-platform roadmap. He became valued for ruthless backlog clarity, calm stakeholder alignment, and an unnerving ability to turn vague requests into testable outcomes.",
        "He persuaded facilities, workplace technology, and leadership to sponsor an after-hours golf outing around a supposedly harmless 'course optimization pilot'—a sensor-and-operations product experiment measuring turf consistency, pace of play, and employee engagement.",
        "When the pilot produces irregular data and the course fails his standards, Joe refuses to close the increment. He commandeers a mower and the maintenance controls to complete the sprint goal before morning.",
        "The player is a coworker left behind when the gates lock. To Joe, they are no longer a colleague; they are unplanned scope, an unmanaged dependency, and the last blocker between him and release acceptance.",
        "Joe's Product Owner habits make him frighteningly methodical beneath his erratic movements: he observes behavior, reorders priorities, tests hypotheses, clarifies acceptance criteria, and removes blockers.",
        "His mower stripes are a physical roadmap. Each cut lane is an accepted increment; each patch of rough is technical debt; every footprint is a defect report; every hiding place is a dependency he has not resolved.",
        "Joe sincerely believes he is maximizing value for the course, the company, and every stakeholder. The horror deepens because he never sees himself as angry or cruel—only accountable for the product outcome.",
        "Keep Joe sincere and competent. The satire belongs in the collision between corporate process, golf etiquette, and the increasingly impossible course.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Unresolved horror questions", 2)
    for item in (
        "Why do the course's gates, sprinklers, speakers, and scoreboards recognize Joe's insurance-company credentials?",
        "Why does a weathered scorecard already contain the player's employee number?",
        "How does Joe's optimization deck predict routes the player has not taken yet?",
        "Why does the clubhouse clock reset whenever the player reaches the ninth hole?",
    ):
        add_bullet(doc, item)


def add_hole_progression(doc):
    add_heading(doc, "6. Nine-hole campaign progression", 1)
    add_data_table(
        doc,
        ["Hole", "New mechanic", "Horror or dramatic escalation", "Comedy or release"],
        [
            ("1", "Basic concealment and straight mower charge", "Joe is introduced at a long, readable distance.", "An etiquette sign explains how to let him play through."),
            ("2", "Persistent mowing and shrinking cover", "A previously safe route is gone on the return trip.", "A scorecard calls hiding in rough a pace-of-play violation."),
            ("3", "Sprinklers and sound masking", "Audio direction becomes unreliable during irrigation cycles.", "Joe stops mid-chase to inspect a sprinkler head."),
            ("4", "Golf cart traversal", "Speed becomes tempting but attracts Joe from across the course.", "The cart's reverse alarm refuses to stop."),
            ("5", "Bunkers, ponds, and footprints", "Joe predicts exits from visible sand tracks.", "A rake must be returned to its designated place to unlock a shed."),
            ("6", "Silent search without mower", "The engine stops and Joe begins appearing inside structures.", "The unattended mower continues following a perfect stripe."),
            ("7", "Course geometry shifts", "Paths and signs disagree with the player's map.", "Every sign still insists the clubhouse is 300 yards away."),
            ("8", "False mower audio and shortcut use", "Multiple engine sources pressure the player into bad routes.", "The PA congratulates Joe for excellent course maintenance."),
            ("9", "Nearly all rough has been removed", "The final crossing is exposed, fast, and visually simple.", "The exit gate requests one final scorecard signature."),
        ],
        [620, 2080, 3760, 2900],
    )


def add_vertical_slice(doc):
    add_heading(doc, "7. Vertical slice definition", 1)
    add_callout(
        doc,
        "Goal",
        "Prove that listening for Joe, manipulating sound, hiding in persistent grass, and surviving a momentum-based charge are enjoyable inside one 10-15 minute hole.",
        fill=LIGHT_GREEN,
    )
    add_heading(doc, "Required slice content", 2)
    required = (
        "A concise opening establishes the insurance-company golf outing, Joe's role as product owner, and the player as a stranded coworker.",
        "One fairway, one rough field, one bunker, one pond, one maintenance shed, and one tee checkpoint.",
        "A two-part objective: restore sprinkler pressure, then obtain the exit key.",
        "Golf-ball throwing, crouch, sprint, stamina, and one sprinkler control.",
        "A 14-second Delivery Chain that scores linked progress, ball recoveries, bunker baits, optional Change Requests, and contact breaks without modifying survival difficulty.",
        "Persistent grass removal and one opened maintenance shortcut.",
        "Joe states: mowing, listening, investigating, watching, and charging.",
        "One optional par-three shot with a small reward.",
        "One short golf-cart escape or scripted end beat.",
        "Functional restart, checkpoint, settings, and readable failure feedback.",
    )
    for item in required:
        add_bullet(doc, item)

    add_heading(doc, "Explicitly out of scope", 2)
    for item in (
        "All nine holes, complete narrative, or final visual polish.",
        "Combat, weapon upgrades, or a conventional boss fight.",
        "Full procedural generation.",
        "Multiplayer or online services.",
        "Complex inventory crafting.",
        "Production-ready Joe animation in every direction.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Acceptance criteria", 2)
    criteria = (
        "A new player can explain why Joe detected them after one or two failures.",
        "At least two terrain-based escape choices are viable during a charge.",
        "Persistent mowing changes the preferred route during the same attempt.",
        "The mower's direction and aggression are understandable without a minimap.",
        "The optional golf shot creates a meaningful risk rather than a detached minigame.",
        "A player can explain what extended or ended a Delivery Chain and can ignore it without losing access to either exit.",
        "The slice produces at least one genuine scare and one unforced laugh.",
    )
    for criterion in criteria:
        add_bullet(doc, criterion)


def add_technical_standup(doc):
    add_heading(doc, "8. Technical stand-up", 1)
    add_heading(doc, "Recommended architecture", 2)
    for item in (
        "Godot 3D world with billboarded Joe sprites and low-resolution render scaling.",
        "Data-driven Joe state machine with explicit perception, navigation, and animation layers.",
        "Noise-event service that records source position, loudness, category, and decay.",
        "Persistent course-state service for grass, gates, sprinklers, props, and checkpoints.",
        "Course segment scenes that can be tested independently and assembled into holes.",
        "Animation metadata that maps behavioral states to Joe's sprite clips and telegraphs.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Directory structure", 2)
    directory_text = (
        "JoeMowerHorror/\n"
        "  README.md\n"
        "  docs/                 project-facing documents\n"
        "  design/               mechanics, levels, narrative, tuning\n"
        "  game/                 engine project and runtime source\n"
        "  assets/\n"
        "    characters/joe/      current Joe sprite package\n"
        "    environment/         terrain, props, foliage, structures\n"
        "    audio/               mower, ambience, UI, PA, scares\n"
        "  builds/                local packaged builds\n"
        "  qa/                    review output and test notes\n"
        "  tools/                 project-local content utilities"
    )
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    for index, line in enumerate(directory_text.splitlines()):
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9.2, color=INK)
        if index < len(directory_text.splitlines()) - 1:
            run.add_break()

    add_heading(doc, "Initial asset backlog", 2)
    backlog = (
        "Joe: idle/listen, mowing patrol, charge, collision recovery, silent-search walk, distant watch, and fail-state animation.",
        "Course: fairway, rough, cut-grass states, bunker, pond, cart path, tree line, fog, and sprinkler particles.",
        "Props: mower, golf balls, club, cart, rake, flags, signs, gates, keys, valves, shed, PA speakers, and vending machine.",
        "Audio: six mower states, pull-start, shutdown, foot movement, grass, golf impacts, sprinklers, cart, PA, and stingers.",
        "UI: scorecard, interaction prompt, stamina treatment, objective list, settings, pause, and failure stamp.",
    )
    for item in backlog:
        add_bullet(doc, item)


def add_plan_and_risks(doc):
    add_heading(doc, "9. Production plan", 1)
    add_data_table(
        doc,
        ["Phase", "Duration", "Primary output", "Exit condition"],
        [
            ("Preproduction", "1 week", "Movement test, art-scale test, Joe cue language, one graybox hole.", "The camera, sprite scale, and course readability feel coherent."),
            ("Systems prototype", "2 weeks", "Noise events, Joe state machine, grass concealment, mowing persistence.", "Joe can detect, investigate, mow, and charge with understandable causes."),
            ("Vertical-slice content", "2 weeks", "Objectives, tools, terrain interactions, golf challenge, checkpoint.", "The complete 10-15 minute loop is playable."),
            ("Presentation pass", "1-2 weeks", "Lighting, fog, mower audio, PA comedy, UI, failure and end beats.", "The slice produces intended suspense, scare, and release beats."),
            ("Playtest and tuning", "1 week", "Telemetry notes, route tuning, cue timing, difficulty options.", "Most players understand detection and can recover from at least one mistake."),
        ],
        [1500, 1350, 3650, 2860],
    )

    add_heading(doc, "Primary design risks", 2)
    risks = [
        ("Joe feels random instead of readable", "Tie visible head, hand, posture, and throttle cues directly to state transitions."),
        ("Constant mower audio becomes tiring", "Use distance, occlusion, silence, false cues, and short pursuit windows."),
        ("Billboard presentation breaks at close range", "Control minimum distance, use angle variants, and trigger authored close-contact framing."),
        ("Comedy weakens the threat", "Keep Joe sincere and place jokes after tension peaks or inside neutral course systems."),
        ("Grass interaction is too expensive", "Use coarse gameplay cells and visual shader responses instead of simulating every blade."),
        ("Golf challenge feels disconnected", "Make shots noisy, route-dependent, and tied to useful rewards."),
    ]
    for label, mitigation in risks:
        add_body(doc, f"{label}. {mitigation}", bold_lead=f"{label}.")


def add_names(doc):
    add_heading(doc, "10. Title direction and alternatives", 1)
    add_callout(
        doc,
        "Selected working title",
        "Rough Cut: A Joe Horror Game. The name keeps Joe at the center while the title connects tall grass, mower blades, a violent threat, and the product-owner idea of an unfinished release. It is short, distinctive, and supports both dread and dry comedy.",
    )
    add_data_table(
        doc,
        ["Name", "Tone", "Why it works"],
        [
            ("Rough Cut", "Selected: stylish horror", "Connects grass, mowing, editing, product releases, and violence without being graphic."),
            ("The Back Nine", "Balanced horror", "Familiar golf language that also suggests a hidden or final ordeal."),
            ("Please Play Through", "Deadpan horror-comedy", "Sounds polite while implying that something dangerous is behind you."),
            ("Groundskeeper Joe", "Character-forward", "Simple, memorable, and makes Joe the franchise identity."),
            ("Cart Path Only", "Deadpan comedy", "A mundane rule that becomes threatening when Joe enforces it."),
            ("No Mulligans", "Horror-comedy", "Communicates consequence and replay without naming the monster."),
            ("The Course Is Closed", "Atmospheric horror", "Immediately establishes trespass, isolation, and after-hours danger."),
            ("Last Mowing", "Creature-feature", "Directly communicates the antagonist's tool and an end-of-night event."),
            ("Cut Below Par", "Pun-forward", "Combines mowing and golf performance with a dark edge."),
            ("Stay on the Fairway", "Threatening instruction", "Works as both course advice and Joe's controlling demand."),
            ("Joe's Final Round", "Character and golf", "Supports a narrative campaign built around nine holes."),
            ("Tee Time with Joe", "Comedy-forward", "Friendly phrasing that becomes unsettling through context."),
            ("After Closing", "Broad horror", "Strong atmosphere but less distinctive without the subtitle or key art."),
            ("Fore!", "Minimal and punchy", "Recognizable warning language with strong marketing potential."),
            ("Par for the Corpse", "Camp horror-comedy", "Memorable and explicit, but commits the project to a sillier tone."),
        ],
        [2450, 1600, 5310],
    )
    add_body(
        doc,
        "Name availability, trademarks, domains, and storefront conflicts have not been researched. Perform a clearance check before public announcement or commercial use.",
    )

    add_heading(doc, "Tagline ideas", 2)
    for item in (
        "The course closes at dusk. Joe does not.",
        "Every blade is in scope.",
        "Your coverage ends at dusk.",
        "Please allow faster groundskeepers to play through.",
        "Nine holes. One mower. No mulligans.",
        "Stay on the fairway.",
        "He keeps the course immaculate.",
    ):
        add_bullet(doc, item)


def add_next_steps(doc):
    add_heading(doc, "11. Immediate next steps", 1)
    for step in (
        "Use Rough Cut as the internal working title; complete trademark, domain, and storefront clearance before a public announcement.",
        "Create the Godot project inside the existing game directory.",
        "Build a low-resolution rendering and billboard-sprite proof before producing more art.",
        "Graybox one hole with fairway, rough, bunker, pond, and maintenance shed.",
        "Implement noise events and Joe's mowing, listening, investigating, and charging states.",
        "Add persistent grass removal and one sprinkler interaction.",
        "Playtest the chase before adding narrative, comedy, or additional holes.",
        "Record a 60-second internal gameplay capture and evaluate readability at native resolution.",
    ):
        add_number(doc, step)

    add_callout(
        doc,
        "First decision",
        "Do not begin with nine holes. Prove one memorable pursuit in one small course segment, then expand only after Joe is readable, frightening, and fun to manipulate.",
        fill=LIGHT_GREEN,
    )


def build_document():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    properties = doc.core_properties
    properties.title = "Rough Cut Game Blueprint"
    properties.subject = "Game concept, vertical slice, art direction, production plan, and naming"
    properties.author = "Rough Cut Project"
    properties.keywords = "Rough Cut, horror game, golf course, insurance, product owner, lawn mower, Joe, pixel art, game design"

    add_cover(doc)
    add_snapshot(doc)
    add_format_and_loop(doc)
    add_joe_system(doc)
    add_player_course(doc)
    add_tone_audio(doc)
    add_hole_progression(doc)
    add_vertical_slice(doc)
    add_technical_standup(doc)
    add_plan_and_risks(doc)
    add_names(doc)
    add_next_steps(doc)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
